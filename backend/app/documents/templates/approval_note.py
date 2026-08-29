"""Approval Note Word Document (DOCX) Template Renderer (TRD Section 22, PRD Requirement #8)."""

import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_background(cell, fill_hex: str):
    """Set background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner padding of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def render_approval_note(
    data: Dict[str, Any],
    output_path: Path,
) -> Path:
    """
    Render a professional Sovereign AI Approval Note (.docx).
    Contains all 7 mandatory sections:
    1. Title Block
    2. Inspection Summary
    3. Critical Findings
    4. Compliance Gaps & SOP Citations
    5. Actionable Recommendations
    6. AI-Generated Draft Disclaimer & Footer
    7. Exact Generation Timestamp
    """
    doc = docx.Document()

    # Set document margins (0.75 in)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        # Header & Footer with AI-Generated disclaimer (PRD Req #8)
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("CONFIDENTIAL / SOVEREIGN ON-PREMISE AI WORKBENCH")
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(128, 128, 128)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("AI-Generated Draft — For Internal Engineering Review Only (Non-Certified Verdict)")
        frun.font.size = Pt(8.5)
        frun.font.italic = True
        frun.font.color.rgb = RGBColor(180, 50, 50)

    # 1. Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(data.get("title", "TECHNICAL APPROVAL NOTE: DOCUMENT COMPLIANCE REVIEW"))
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(16, 44, 87)  # Deep Navy

    # Metadata Subtitle Box
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    meta_items = [
        ("Task ID / Ref:", str(data.get("task_id", "N/A"))),
        ("Facility / Source:", str(data.get("facility", f"Source: {data.get('source_document', 'Uploaded Document')}"))),
        ("Generated Timestamp (UTC):", str(data.get("timestamp", datetime.datetime.now(datetime.timezone.utc).isoformat()))),
        ("Evaluation Status:", str(data.get("status", "Analyzed via Local Model (On-Premise)"))),
    ]

    for idx, (label, val) in enumerate(meta_items):
        row = meta_table.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.2)
        c1.width = Inches(4.8)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = RGBColor(50, 50, 50)

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        if "ACTION REQUIRED" in val or "CRITICAL" in val:
            r1.font.bold = True
            r1.font.color.rgb = RGBColor(180, 0, 0)

        set_cell_background(c0, "F0F4F8")
        set_cell_background(c1, "FAFAFA")
        set_cell_margins(c0, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)

    doc.add_paragraph()  # Spacing

    # 2. Inspection Summary Section
    h1 = doc.add_paragraph()
    r = h1.add_run("1. Inspection Overview & Summary")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(16, 44, 87)

    summary_text = data.get(
        "summary",
        "Document intelligence analysis was executed on the submitted source document and verified against relevant standards."
    )
    p_sum = doc.add_paragraph(summary_text)
    p_sum.style.font.size = Pt(10)

    # 3. Critical Findings Section
    h2 = doc.add_paragraph()
    r = h2.add_run("2. Critical Inspection Findings")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(16, 44, 87)

    findings = data.get("critical_findings") or [
        "No critical findings identified in the analyzed document."
    ]
    for f in findings:
        p_f = doc.add_paragraph(style="List Bullet")
        rf = p_f.add_run(f)
        rf.font.size = Pt(10)

    # 4. Compliance Gaps & SOP Citations Section
    h3 = doc.add_paragraph()
    r = h3.add_run("3. Compliance Gaps & Authoritative SOP Citations")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(16, 44, 87)

    citations = data.get("citations") or []
    gaps = data.get("compliance_gaps") or [
        ("General Inspection Scope", "Standard Operating Baseline", "REVIEW COMPLETED")
    ]

    gap_table = doc.add_table(rows=1, cols=3)
    gap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    gap_table.autofit = False

    headers = ["Observed Finding / Defect", "Authoritative SOP Citation", "Compliance Rating"]
    hdr_row = gap_table.rows[0]
    col_widths = [Inches(2.5), Inches(3.0), Inches(1.5)]
    for i, h_text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = col_widths[i]
        hp = cell.paragraphs[0]
        hr = hp.add_run(h_text)
        hr.font.bold = True
        hr.font.size = Pt(9.5)
        hr.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "102C57")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

    for item in gaps:
        row = gap_table.add_row()
        for col_idx, text_val in enumerate(item):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            p = cell.paragraphs[0]
            r = p.add_run(text_val)
            r.font.size = Pt(9)
            if "CRITICAL" in text_val or "NON-COMPLIANCE" in text_val:
                r.font.bold = True
                r.font.color.rgb = RGBColor(180, 0, 0)
            set_cell_background(cell, "F9F9F9")
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    doc.add_paragraph()

    # 5. Actionable Recommendations Section
    h4 = doc.add_paragraph()
    r = h4.add_run("4. Actionable Engineering Recommendations")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(16, 44, 87)

    recs = data.get("recommendations") or [
        "Maintain routine monitoring as per standard operational guidelines."
    ]
    for idx, rec in enumerate(recs, 1):
        p_r = doc.add_paragraph()
        r_num = p_r.add_run(f"4.{idx} ")
        r_num.font.bold = True
        r_num.font.size = Pt(10)
        r_body = p_r.add_run(rec)
        r_body.font.size = Pt(10)

    # 6. Sign-off & Disclaimer Block
    doc.add_paragraph()
    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.autofit = False

    s0, s1 = sign_table.rows[0].cells[0], sign_table.rows[0].cells[1]
    s0.paragraphs[0].add_run("Prepared Autonomously By:\nSovereign AI Workbench (General Model Agent)").font.size = Pt(9)
    s1.paragraphs[0].add_run("Reviewed & Endorsed By:\n____________________________\nLead Mechanical Maintenance Engineer").font.size = Pt(9)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
