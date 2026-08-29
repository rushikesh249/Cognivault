"""Structured Analysis Report DOCX Template Renderer (TRD Section 22, ADR-007).

Renders a grounded document-analysis deliverable (e.g. Research Paper
Analysis): Main Topic, Objectives, Methodology, Key Findings, Conclusions,
Overall Summary, and a Sources & Grounding section identifying exactly which
document and knowledge-base chunks were used.
"""

import datetime
from pathlib import Path
from typing import Any, Dict, List
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from backend.app.documents.templates.approval_note import set_cell_background, set_cell_margins


def render_structured_report(
    data: Dict[str, Any],
    output_path: Path,
) -> Path:
    """
    Render a source-grounded structured analysis report (.docx).

    Expected payload keys:
    - title, task_id, source_document, timestamp, status
    - summary: concise overall summary
    - sections: list of {"heading": str, "content": str | list[str]}
    - sources: list of source/citation strings used for grounding
    """
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("CONFIDENTIAL / SOVEREIGN ON-PREMISE AI WORKBENCH")
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(128, 128, 128)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("AI-Generated Draft — For Internal Review Only (Non-Certified)")
        frun.font.size = Pt(8.5)
        frun.font.italic = True
        frun.font.color.rgb = RGBColor(180, 50, 50)

    # 1. Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(str(data.get("title", "Document Analysis Report")))
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(16, 44, 87)  # Deep Navy

    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    meta_items = [
        ("Task ID / Ref:", str(data.get("task_id", "N/A"))),
        ("Source Document:", str(data.get("source_document", "Not specified"))),
        ("Generated Timestamp (UTC):", str(data.get("timestamp", datetime.datetime.now(datetime.timezone.utc).isoformat()))),
        ("Analysis Engine:", str(data.get("status", "Local General Model (On-Premise)"))),
    ]

    for idx, (label, val) in enumerate(meta_items):
        row = meta_table.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.2)
        c1.width = Inches(4.8)

        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = RGBColor(50, 50, 50)

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)

        set_cell_background(c0, "F0F4F8")
        set_cell_background(c1, "FAFAFA")
        set_cell_margins(c0, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)

    doc.add_paragraph()  # Spacing

    # 2. Analysis Sections (Main Topic, Objectives, Methodology, ...)
    sections_payload: List[Dict[str, Any]] = data.get("sections", [])
    for sec_idx, sec in enumerate(sections_payload, start=1):
        heading = str(sec.get("heading", f"Section {sec_idx}"))
        content = sec.get("content", "")

        h = doc.add_paragraph()
        r = h.add_run(f"{sec_idx}. {heading}")
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(16, 44, 87)

        if isinstance(content, list):
            for item in content:
                p_b = doc.add_paragraph(style="List Bullet")
                r_b = p_b.add_run(str(item))
                r_b.font.size = Pt(10)
        else:
            p_c = doc.add_paragraph(str(content))
            p_c.style.font.size = Pt(10)

    # 3. Sources & Grounding Section
    h_src = doc.add_paragraph()
    r_src = h_src.add_run(f"{len(sections_payload) + 1}. Sources & Grounding")
    r_src.font.size = Pt(13)
    r_src.font.bold = True
    r_src.font.color.rgb = RGBColor(16, 44, 87)

    sources: List[str] = data.get("sources", []) or data.get("citations", [])
    if sources:
        for src in sources:
            p_s = doc.add_paragraph(style="List Bullet")
            r_s = p_s.add_run(str(src))
            r_s.font.size = Pt(10)
    else:
        p_none = doc.add_paragraph(
            "No external knowledge-base sources were used. This report is grounded "
            "exclusively in the uploaded source document."
        )
        p_none.style.font.size = Pt(10)

    grounding_note = data.get("grounding_note")
    if grounding_note:
        p_g = doc.add_paragraph(str(grounding_note))
        p_g.style.font.size = Pt(9)
        p_g.runs[0].font.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
