"""Hero Flow 1 End-to-End Integration Test (TRD Section 19, Table 44, Test Plan Section 12)."""

import uuid
from pathlib import Path
import docx
import pytest
from fastapi.testclient import TestClient

from backend.app.agent.graph import agent_graph
from backend.app.agent.state import AgentState
from backend.app.main import app
from backend.app.persistence.artifact_repository import ArtifactRepository
from backend.app.persistence.db import get_db_context, init_db
from backend.app.persistence.file_repository import FileRepository
from backend.app.persistence.task_repository import TaskRepository


@pytest.fixture(autouse=True)
def setup_db():
    init_db()


def test_hero_flow_1_document_intelligence_end_to_end():
    """
    Hero Flow 1 (TRD Section 19, Table 44):
    Upload scanned PDF -> OCR -> RAG search -> LLM comparison -> DOCX generation -> Validation -> Artifact storage.
    """
    client = TestClient(app)
    goal_prompt = "Analyze uploaded equipment inspection report MRPL-INSP-2026-8842 against safety SOPs and generate technical Approval Note."

    # 1. Create task in SQLite
    with get_db_context() as session:
        task_repo = TaskRepository(session)
        task = task_repo.create(
            title="Hero Flow 1: Refinery Inspection Compliance Review",
            task_type="document",
            prompt=goal_prompt,
        )
        task_id = task.task_id

    # 2. Attach synthetic demonstration PDF to task with unique file_id
    demo_pdf = Path("knowledge_base/demo_inputs/scanned_inspection_report.pdf")
    assert demo_pdf.exists(), "Demo inspection report PDF must exist"

    unique_file_id = f"demo-{uuid.uuid4().hex[:8]}"
    with get_db_context() as session:
        file_repo = FileRepository(session)
        file_repo.create(
            file_id=unique_file_id,
            task_id=task_id,
            filename="scanned_inspection_report.pdf",
            mime_type="application/pdf",
            pages=2,
            size_bytes=demo_pdf.stat().st_size,
            storage_path=str(demo_pdf.resolve()),
        )

    # 3. Initialize Agent State and execute LangGraph workflow
    initial_state: AgentState = {
        "task_id": task_id,
        "task_type": "document",
        "goal": goal_prompt,
        "plan": [],
        "current_step_index": 0,
        "iteration": 0,
        "max_iterations": 4,
        "selected_model_id": None,
        "tool_calls": [],
        "observations": [],
        "validation_passed": False,
        "validation_notes": None,
        "final_artifact_id": None,
        "status": "running",
        "error": None,
        "_staged_tool_call": None,
        "_raw_execution_result": None,
    }

    final_state = agent_graph.invoke(initial_state, config={"recursion_limit": 100})

    # 4. State & Terminal Assertions
    assert final_state["status"] == "succeeded"
    assert final_state["validation_passed"] is True
    assert final_state["iteration"] == 1
    artifact_id = final_state.get("final_artifact_id")
    assert artifact_id is not None, "Hero Flow 1 must produce a final_artifact_id"

    # 5. Database Assertions
    with get_db_context() as session:
        art_repo = ArtifactRepository(session)
        art = art_repo.get_by_id(artifact_id)
        assert art is not None
        assert art.task_id == task_id
        assert art.kind == "docx"
        output_file_path = Path(art.storage_path)

    assert output_file_path.exists(), f"Rendered artifact file missing at: {output_file_path}"

    # 6. Artifact Retrieval API Tests (TRD Table 18)
    # 6a. Binary Download
    res_bin = client.get(f"/api/artifacts/{artifact_id}")
    assert res_bin.status_code == 200
    assert len(res_bin.content) > 0
    assert "wordprocessingml" in res_bin.headers["content-type"]

    # 6b. JSON Metadata
    res_meta = client.get(f"/api/artifacts/{artifact_id}?meta=true")
    assert res_meta.status_code == 200
    meta_data = res_meta.json()
    assert meta_data["artifact_id"] == artifact_id
    assert meta_data["kind"] == "docx"
    assert len(meta_data["sources"]) > 0

    # 7. Document Content & Section Integrity Verification via python-docx
    doc = docx.Document(str(output_file_path))
    doc_text = "\n".join(p.text for p in doc.paragraphs)
    
    assert "APPROVAL NOTE" in doc_text.upper()
    assert "Inspection Overview" in doc_text
    assert "Critical Inspection Findings" in doc_text
    assert "Compliance Gaps" in doc_text
    assert "Actionable Engineering Recommendations" in doc_text

    # Verify AI disclaimer in footer
    footer_text = "\n".join(p.text for s in doc.sections for p in s.footer.paragraphs)
    assert "AI-Generated Draft" in footer_text
