"""Automated Regression Test Suite for Vision Workflow & Artifact Grounding (Item #10)."""

import io
import json
import uuid
import docx
import pytest
from pathlib import Path
from PIL import Image

from backend.app.core.config import settings
from backend.app.models.model_registry import ModelRegistry
from backend.app.models.provider import ModelProvider
from backend.app.models.router import ModelRouter
from backend.app.models.schema import ModelConfig, ModelStatus, ProviderStatus
from backend.app.multimodal.vision_service import VisionResult, VisionService
from backend.app.persistence.artifact_repository import ArtifactRepository
from backend.app.persistence.db import get_db_context, init_db
from backend.app.persistence.file_repository import FileRepository
from backend.app.persistence.task_repository import TaskRepository
from backend.app.services.agent_service import AgentService
from backend.app.services.vision_service import VisionAppService, set_vision_app_service


class RealisticFlangeInspectionVisionProvider(ModelProvider):
    """Local VLM stub providing realistic industrial flange inspection observations."""

    def is_provider_available(self) -> bool:
        return True

    def is_model_available(self, model_identifier: str) -> bool:
        return True

    def get_provider_status(self) -> ProviderStatus:
        return ProviderStatus.AVAILABLE

    def get_model_status(self, model_config: ModelConfig) -> ModelStatus:
        return ModelStatus.AVAILABLE

    def list_available_models(self):
        return ["llava:7b-v1.5-q4_K_M"]

    def ensure_loaded(self, model_id: str) -> bool:
        return True

    def unload(self, model_id: str) -> bool:
        return True

    def unload_lru(self) -> bool:
        return True

    def generate(self, model_id: str, prompt: str, images=None, system=None, format=None, stream=False) -> str:
        return json.dumps({
            "observation": [
                "Industrial pipe/flange assembly with bolted connection",
                "Visible surface corrosion and rust around the joint seam",
                "Peeling and deteriorated protective paint coating on lower quadrant",
                "Visible hex head screws and fasteners securing flange perimeter",
            ],
            "interpretation": [
                "Local atmospheric moisture accumulation causing progressive oxide layer build-up",
            ],
            "uncertainty": [
                "Camera angle and resolution preclude internal wall thickness measurement",
            ],
        })


def test_vision_workflow_and_artifact_grounding_regression():
    """
    Verify requirements:
    1. Vision task selects local-vision-model
    2. VLM observations are generated
    3. Observations reach final_deliverable
    4. Final artifact contains 'corrosion' and 'peeling paint'
    5. Final artifact does NOT incorrectly report 'Not found' for key visual findings
    6. Final artifact metadata does NOT say 'local-general-model' when local-vision-model was selected
    """
    init_db()
    task_id = f"vision-grounding-{uuid.uuid4()}"
    file_id = f"file-{uuid.uuid4()}"

    # 1. Setup synthetic image
    demo_inputs = Path("knowledge_base/demo_inputs")
    demo_inputs.mkdir(parents=True, exist_ok=True)
    img_path = demo_inputs / f"{file_id}.jpg"
    img = Image.new("RGB", (256, 256), color=(120, 80, 50))
    img.save(img_path, format="JPEG")

    # 2. Register task and file
    with get_db_context() as session:
        t_repo = TaskRepository(session)
        f_repo = FileRepository(session)

        t_repo.create(
            task_id=task_id,
            title="Heat Exchanger Pipe Flange Inspection",
            task_type="vision",
            prompt="Analyze industrial flange joint for visible surface corrosion, paint condition, and fastener integrity.",
        )
        f_repo.create(
            file_id=file_id,
            task_id=task_id,
            filename="flange_inspection.jpg",
            mime_type="image/jpeg",
            size_bytes=img_path.stat().st_size,
            storage_path=str(img_path),
        )

    # 3. Wire realistic vision provider
    stub_provider = RealisticFlangeInspectionVisionProvider()
    vision_domain = VisionService(provider=stub_provider)
    set_vision_app_service(VisionAppService(domain_service=vision_domain))

    # 4. Execute workflow
    agent_svc = AgentService()
    final_state = agent_svc._run_graph_sync(task_id)

    # Assertions on state and execution
    assert final_state["status"] == "succeeded"
    assert final_state["validation_passed"] is True
    assert final_state["selected_model_id"] == "local-vision-model"

    artifact_id = final_state.get("final_artifact_id")
    assert artifact_id is not None

    # Verify deliverable artifact content
    with get_db_context() as session:
        art_repo = ArtifactRepository(session)
        art = art_repo.get_by_id(artifact_id)
        assert art is not None
        assert Path(art.storage_path).exists()

        # Read DOCX document text
        doc = docx.Document(art.storage_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        full_text_lower = full_text.lower()

        # Check metadata in tables
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "

        combined_doc = (full_text + " " + table_text).lower()

        # 1. Check model identity is local-vision-model
        assert "local-general-model" not in combined_doc
        assert "local-vision-model" in combined_doc

        # 2. Check key visual observations are present
        assert "corrosion" in combined_doc
        assert "peeling" in combined_doc
        assert "flange" in combined_doc
        assert "fastener" in combined_doc or "screws" in combined_doc

        # 3. Check NO 'Not found' or generic document analysis filler
        assert "not found" not in combined_doc
        assert "methodology" not in combined_doc
        assert "objectives" not in combined_doc

        # 4. Check Stage 4 tool_selection event exists
        t_repo = TaskRepository(session)
        events = t_repo.get_events(task_id)
        node_names = [e.node for e in events]
        assert "tool_selection" in node_names
        assert any(e.node == "tool_selection" and ("vision analysis tool" in e.message or "report generation tool" in e.message) for e in events)


class IndustrialWorkshopVisionProvider(ModelProvider):
    """Local VLM stub providing realistic observations of industrial workshop with site conditions."""

    def is_provider_available(self) -> bool:
        return True

    def is_model_available(self, model_identifier: str) -> bool:
        return True

    def get_provider_status(self) -> ProviderStatus:
        return ProviderStatus.AVAILABLE

    def get_model_status(self, model_config: ModelConfig) -> ModelStatus:
        return ModelStatus.AVAILABLE

    def list_available_models(self):
        return ["llava:7b-v1.5-q4_K_M"]

    def ensure_loaded(self, model_id: str) -> bool:
        return True

    def unload(self, model_id: str) -> bool:
        return True

    def unload_lru(self) -> bool:
        return True

    def generate(self, model_id: str, prompt: str, images=None, system=None, format=None, stream=False) -> str:
        return json.dumps({
            "observation": [
                "Industrial workshop area containing heavy processing machinery and piping runs",
                "Muddy ground and standing water puddle visible near the equipment foundation",
                "Exposed exterior brick wall and stacked cinder blocks along the perimeter",
                "Flanged steel pipe connection with visible secure fastening hardware",
            ],
            "interpretation": [
                "Machinery and piping located in semi-exposed outdoor or workshop facility",
            ],
            "uncertainty": [
                "Operational integrity, internal pressure, and subsurface wear cannot be verified from photograph alone",
            ],
        })


def test_vision_docx_terminology_and_site_conditions_regression():
    """
    Regression test verifying:
    1. Stage 4 tool_selection event exists in task events.
    2. DOCX terminology uses:
       - Visual Observations
       - Visible Objects / Components
       - Site / Surface Conditions
       - Engineering Interpretation
       - Limitations / Uncertainty
       - Sources & Grounding
    3. Puddles, muddy ground, brick walls, cinder blocks are placed in Site / Surface Conditions, not misclassified as defects.
    4. Unsupported claims (leakage, structural failure, pressure loss, unsafe operation) are NOT generated.
    5. Conservative engineering interpretation is generated dynamically from actual VLM observations.
    6. local-vision-model is used and local-general-model is absent.
    7. Final status is succeeded.
    """
    init_db()
    task_id = f"vision-terminology-{uuid.uuid4()}"
    file_id = f"file-{uuid.uuid4()}"

    demo_inputs = Path("knowledge_base/demo_inputs")
    demo_inputs.mkdir(parents=True, exist_ok=True)
    img_path = demo_inputs / f"{file_id}.jpg"
    img = Image.new("RGB", (256, 256), color=(100, 100, 100))
    img.save(img_path, format="JPEG")

    with get_db_context() as session:
        t_repo = TaskRepository(session)
        f_repo = FileRepository(session)

        t_repo.create(
            task_id=task_id,
            title="Workshop Equipment Area Inspection",
            task_type="vision",
            prompt="Analyze workshop equipment image for components, site surface conditions, and structural elements.",
        )
        f_repo.create(
            file_id=file_id,
            task_id=task_id,
            filename="workshop_inspection.jpg",
            mime_type="image/jpeg",
            size_bytes=img_path.stat().st_size,
            storage_path=str(img_path),
        )

    stub_provider = IndustrialWorkshopVisionProvider()
    vision_domain = VisionService(provider=stub_provider)
    set_vision_app_service(VisionAppService(domain_service=vision_domain))

    agent_svc = AgentService()
    final_state = agent_svc._run_graph_sync(task_id)

    # 1. Verification of final status
    assert final_state["status"] == "succeeded"
    assert final_state["validation_passed"] is True
    assert final_state["selected_model_id"] == "local-vision-model"

    artifact_id = final_state.get("final_artifact_id")
    assert artifact_id is not None

    with get_db_context() as session:
        # 2. Verify Stage 4 tool_selection event
        t_repo = TaskRepository(session)
        events = t_repo.get_events(task_id)
        node_names = [e.node for e in events]
        assert "tool_selection" in node_names
        assert any(e.node == "tool_selection" and "vision analysis tool" in e.message for e in events)

        # 3. Verify artifact content & terminology
        art_repo = ArtifactRepository(session)
        art = art_repo.get_by_id(artifact_id)
        assert art is not None
        assert Path(art.storage_path).exists()

        doc = docx.Document(art.storage_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "

        combined_doc = full_text + " " + table_text
        combined_doc_lower = combined_doc.lower()

        # Model validation
        assert "local-general-model" not in combined_doc_lower
        assert "local-vision-model" in combined_doc_lower

        # Exact required headings
        assert "Visual Observations" in combined_doc
        assert "Visible Objects / Components" in combined_doc
        assert "Site / Surface Conditions" in combined_doc
        assert "Engineering Interpretation" in combined_doc
        assert "Limitations / Uncertainty" in combined_doc
        assert "Sources & Grounding" in combined_doc

        # Observations present dynamically
        assert "machinery" in combined_doc_lower or "piping" in combined_doc_lower
        assert "puddle" in combined_doc_lower or "muddy" in combined_doc_lower
        assert "brick" in combined_doc_lower or "cinder block" in combined_doc_lower

        # Verify ordinary items are NOT called defects
        assert "defects observed" not in combined_doc_lower
        assert "defect" not in combined_doc_lower

        # Verify unsupported catastrophic claims are NOT present
        assert "structural failure" not in combined_doc_lower
        assert "pressure loss" not in combined_doc_lower
        assert "leakage" not in combined_doc_lower
        assert "equipment failure" not in combined_doc_lower
        assert "unsafe operation" not in combined_doc_lower

        # Conservative engineering interpretation verified
        assert "operational integrity" in combined_doc_lower or "cannot be confirmed" in combined_doc_lower

