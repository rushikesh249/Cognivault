"""Regression tests for uploaded-document grounding in the Document RAG pipeline.

Guards against the historical defect where a "Research Paper Analysis" task on
an uploaded research paper produced a canned "TECHNICAL APPROVAL NOTE:
EQUIPMENT INSPECTION COMPLIANCE" DOCX built from unrelated demo knowledge-base
content instead of the uploaded document.

Covered requirements:
- Uploaded document is the primary source (RC: file linkage + full extraction).
- Extracted text, analysis context, and final DOCX all carry the document's
  distinctive facts.
- No unrelated industrial demo content may appear in the generated report.
- Source/document metadata is preserved (artifact sources attribution).
"""

import json
from pathlib import Path
import docx
import fitz  # PyMuPDF
import pytest

from backend.app.agent.graph import agent_graph
from backend.app.agent.state import AgentState
from backend.app.documents.templates.structured_report import render_structured_report
from backend.app.persistence.artifact_repository import ArtifactRepository
from backend.app.persistence.db import get_db_context, init_db
from backend.app.persistence.file_repository import FileRepository
from backend.app.services.task_service import TaskService

SYNTHETIC_PAPER_TEXT = """NeuroGrid-7T: A Low-Power Spiking Neural Architecture for Edge Vision

Abstract
This paper presents NeuroGrid-7T, a spiking neural architecture that reduces
edge inference energy consumption by 42.7 percent on the Zebrafish-9
benchmark corpus.

Objectives
The primary objective is to reduce inference energy by 42.7 percent while
maintaining accuracy above 85 percent on embedded vision workloads.

Methodology
We trained NeuroGrid-7T on the Zebrafish-9 benchmark corpus using
event-driven sparse activation and evaluated it on Kalpana-class edge
devices with 2W power budgets.

Key Findings
NeuroGrid-7T achieves 87.3 percent top-1 accuracy. Energy consumption
dropped by 42.7 percent compared to the baseline convolutional model.

Conclusions
NeuroGrid-7T is deployable on Kalpana-class edge devices and meets the
2W power budget.
"""

DISTINCTIVE_FACTS = ["NeuroGrid-7T", "42.7", "Zebrafish-9", "Kalpana-class", "87.3"]

FORBIDDEN_INDUSTRIAL_MARKERS = [
    "TECHNICAL APPROVAL NOTE",
    "PRV-204",
    "Refining Unit 02",
    "P-102A",
    "MRPL-INSP",
    "hydrostatic",
    "Flare Header",
    "mechanical seal",
]

RESEARCH_GOAL = (
    "Analyze the uploaded research paper and produce a concise structured report. "
    "Extract the main topic, objectives, methodology, key findings, and conclusions. "
    "Generate the final report as a DOCX document."
)


@pytest.fixture(autouse=True)
def setup_db():
    init_db()


def _build_synthetic_paper_pdf(path: Path) -> Path:
    """Render a synthetic research paper PDF containing the distinctive facts."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), SYNTHETIC_PAPER_TEXT, fontsize=11)
    doc.save(str(path))
    doc.close()
    return path


def _docx_full_text(path: Path) -> str:
    """Collect all paragraph and table text from a DOCX file."""
    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for section in doc.sections:
        parts.extend(p.text for p in section.header.paragraphs)
        parts.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(parts)


def test_research_paper_analysis_is_grounded_to_uploaded_document(tmp_path):
    """
    End-to-end regression (requirements 10 & 11): upload a synthetic research
    paper, run the full agent graph, and verify the extracted text, analysis
    context, and final DOCX all carry the paper's distinctive facts while no
    unrelated industrial demo content leaks in.
    """
    pdf_path = _build_synthetic_paper_pdf(tmp_path / "synthetic_research_paper.pdf")

    # Simulate the frontend flow: the file is uploaded first (no task yet)...
    with get_db_context() as session:
        file_obj = FileRepository(session).create(
            filename=pdf_path.name,
            mime_type="application/pdf",
            size_bytes=pdf_path.stat().st_size,
            storage_path=str(pdf_path),
            pages=1,
        )
        file_id = file_obj.file_id
        assert file_obj.task_id is None

    # ...then the task is created referencing file_ids (RC-1 linkage fix).
    with get_db_context() as session:
        result = TaskService(session).create_task(
            title="Research Paper Analysis",
            task_type="document",
            prompt=RESEARCH_GOAL,
            file_ids=[file_id],
        )
    task_id = result["task_id"]
    assert result["file_ids"] == [file_id], "Uploaded file must be linked to the task"

    with get_db_context() as session:
        linked = FileRepository(session).get_by_id(file_id)
        assert linked is not None and linked.task_id == task_id

    initial_state: AgentState = {
        "task_id": task_id,
        "task_type": "document",
        "goal": RESEARCH_GOAL,
        "plan": [],
        "current_step_index": 0,
        "iteration": 0,
        "max_iterations": 4,
        "selected_model_id": None,
        "tool_calls": [],
        "observations": [],
        "validation_passed": False,
        "validation_notes": None,
        "final_artifact_id": None,
        "status": "running",
        "error": None,
        "_staged_tool_call": None,
        "_raw_execution_result": None,
    }

    final_state = agent_graph.invoke(initial_state, config={"recursion_limit": 100})

    # Terminal state
    assert final_state["status"] == "succeeded"
    assert final_state["validation_passed"] is True
    artifact_id = final_state.get("final_artifact_id")
    assert artifact_id is not None

    # 1. Extraction observation: the uploaded document's text was extracted.
    extracted_obs = [
        obs for obs in final_state["observations"]
        if obs.get("structured_data", {}).get("tool_name") == "extract_text_from_scan"
        and obs.get("structured_data", {}).get("success")
    ]
    assert extracted_obs, "Extraction step must run against the uploaded document"
    extracted_text = extracted_obs[-1]["structured_data"]["data"]["text"]
    for fact in DISTINCTIVE_FACTS:
        assert fact in extracted_text, f"Extracted text is missing distinctive fact '{fact}'"

    # 2. Analysis observation: grounded structured analysis was produced.
    analysis_obs = [
        obs for obs in final_state["observations"]
        if obs.get("structured_data", {}).get("analysis")
    ]
    assert analysis_obs, "Grounded document analysis step must produce an analysis payload"
    analysis = analysis_obs[-1]["structured_data"]["analysis"]
    assert analysis["source_document"] == pdf_path.name
    analysis_blob = json.dumps(analysis["sections"])
    for fact in DISTINCTIVE_FACTS:
        assert fact in analysis_blob, f"Analysis sections are missing distinctive fact '{fact}'"

    # 3. Generated DOCX: contains requested sections and the paper's facts.
    with get_db_context() as session:
        art = ArtifactRepository(session).get_by_id(artifact_id)
        assert art is not None and art.kind == "docx"
        output_path = Path(art.storage_path)
        sources_json = art.sources_json or "[]"
    assert output_path.exists()

    doc_text = _docx_full_text(output_path)
    for heading in ["Main Topic", "Objectives", "Methodology", "Key Findings", "Conclusions", "Overall Summary", "Sources"]:
        assert heading in doc_text, f"Generated report is missing section '{heading}'"
    for fact in DISTINCTIVE_FACTS:
        assert fact in doc_text, f"Generated report is missing distinctive fact '{fact}'"
    assert pdf_path.name in doc_text, "Report must identify the uploaded source document"

    # 4. No unrelated industrial demo content may appear (requirement 11).
    for marker in FORBIDDEN_INDUSTRIAL_MARKERS:
        assert marker.lower() not in doc_text.lower(), (
            f"Unrelated industrial demo content '{marker}' leaked into the research report"
        )

    # 5. Source metadata preserved: artifact sources attribute the uploaded document.
    sources = json.loads(sources_json)
    assert any("Uploaded document" in s and pdf_path.name in s for s in sources), (
        f"Artifact sources must attribute the uploaded document, got: {sources}"
    )


def test_structured_report_renderer_never_emits_approval_note_content(tmp_path):
    """Renderer-level guarantee: the structured report template cannot produce
    'TECHNICAL APPROVAL NOTE: EQUIPMENT INSPECTION COMPLIANCE' content."""
    output_path = tmp_path / "structured_report.docx"
    render_structured_report(
        {
            "title": "Research Paper Analysis",
            "task_id": "task-xyz",
            "source_document": "synthetic_research_paper.pdf",
            "status": "Analyzed via local-general-model (on-premise)",
            "summary": "The paper introduces NeuroGrid-7T.",
            "sections": [
                {"heading": "Main Topic", "content": "NeuroGrid-7T spiking neural architecture"},
                {"heading": "Objectives", "content": "Reduce inference energy by 42.7 percent"},
                {"heading": "Methodology", "content": "Event-driven sparse activation on Zebrafish-9"},
                {"heading": "Key Findings", "content": ["87.3 percent top-1 accuracy"]},
                {"heading": "Conclusions", "content": "Deployable on Kalpana-class edge devices"},
            ],
            "sources": ["Uploaded document: synthetic_research_paper.pdf"],
        },
        output_path,
    )

    doc_text = _docx_full_text(output_path)
    for marker in FORBIDDEN_INDUSTRIAL_MARKERS:
        assert marker.lower() not in doc_text.lower(), (
            f"Structured report renderer emitted forbidden content '{marker}'"
        )
    assert "EQUIPMENT INSPECTION COMPLIANCE" not in doc_text.upper()
    for fact in DISTINCTIVE_FACTS:
        assert fact in doc_text
